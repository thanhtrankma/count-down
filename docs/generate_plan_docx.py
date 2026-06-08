#!/usr/bin/env python3
"""Generate countdown-project-plan.docx from project plan content."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_prompt_block(doc, title: str, content: str):
    doc.add_heading(title, level=2)
    add_code_block(doc, content.strip())


def main():
    out_path = Path(__file__).parent / "countdown-project-plan.docx"
    doc = Document()

    # Title
    title = doc.add_heading("Countdown Video Generator", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Kế hoạch dự án — Python + FFmpeg + React")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph("Phiên bản: 1.0  |  Ngày: 06/06/2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Overview
    doc.add_heading("1. Tổng quan", level=1)
    doc.add_paragraph(
        "Website cho phép tùy chỉnh countdown (thời gian, style) và export video MP4 dài (vài giờ), "
        "render trên local server bằng FFmpeg (không render frame-by-frame từ browser)."
    )
    add_table(
        doc,
        ["Thành phần", "Công nghệ", "Vai trò"],
        [
            ("Frontend", "Vite + React + TypeScript + Tailwind", "Form cấu hình, preview 10–15s, gọi API, theo dõi progress"),
            ("Backend", "Python 3.11+ + FastAPI", "Nhận config, sinh ASS, gọi FFmpeg, quản lý job"),
            ("Render", "FFmpeg", "Sinh MP4 vài giờ từ ASS + nền tĩnh/gradient"),
            ("Storage", "Local ./output/", "Lưu video + file tạm"),
        ],
    )

    # 2. Architecture
    doc.add_heading("2. Kiến trúc hệ thống", level=1)
    doc.add_heading("Luồng dữ liệu", level=2)
    for line in [
        "Frontend (localhost:5173): Config Form → Preview 10s → Poll Job Progress",
        "Backend (localhost:8000): FastAPI → ASS Generator → Job Manager → FFmpeg Subprocess",
        "Output: MP4 lưu tại ./output/",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Nguyên tắc render", level=2)
    for line in [
        "Mỗi giây countdown = 1 Dialogue event trong file ASS",
        "FFmpeg burn subtitle lên nền solid color",
        "Không dùng Remotion/Puppeteer cho video dài",
        "Web UI chỉ config + preview ngắn; export dài do FFmpeg local",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Thực tế scale (video dài)", level=2)
    add_table(
        doc,
        ["Countdown", "Độ dài video", "Số giây hiển thị", "Frame @ 30fps"],
        [
            ("1 giờ", "~1 giờ", "3.600", "~108.000"),
            ("3 giờ", "~3 giờ", "10.800", "~324.000"),
            ("8 giờ", "~8 giờ", "28.800", "~864.000"),
        ],
    )
    doc.add_paragraph(
        "File 1080p H.264 ~3–5 Mbps: 1 giờ ≈ 1,5–2,5 GB, 8 giờ ≈ 12–20 GB. "
        "FFmpeg + ASS thường render vài phút đến ~30 phút cho clip vài tiếng."
    )

    # 3. Folder structure
    doc.add_heading("3. Cấu trúc thư mục", level=1)
    add_code_block(
        doc,
        """count-down/
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── ConfigForm.tsx
│   │   │   ├── PreviewPlayer.tsx
│   │   │   ├── RenderProgress.tsx
│   │   │   └── FormatSelector.tsx
│   │   ├── hooks/
│   │   │   ├── useCountdown.ts
│   │   │   └── useRenderJob.ts
│   │   ├── types/config.ts
│   │   ├── utils/formatTime.ts
│   │   ├── App.tsx
│   │   └── main.tsx
│   ├── package.json
│   └── vite.config.ts
├── backend/
│   ├── app/
│   │   ├── main.py
│   │   ├── models/schemas.py
│   │   ├── services/
│   │   │   ├── ass_generator.py
│   │   │   ├── ffmpeg_runner.py
│   │   │   └── job_manager.py
│   │   ├── routers/render.py
│   │   └── config.py
│   ├── output/
│   ├── temp/
│   ├── scripts/test_render.py
│   └── requirements.txt
├── shared/config.schema.json
├── .gitignore
└── README.md""",
    )

    # 4. Data contract
    doc.add_heading("4. Data contract (Frontend ↔ Backend)", level=1)
    doc.add_heading("RenderConfig", level=2)
    add_code_block(
        doc,
        """{
  "durationSeconds": 3600,
  "startTime": "03:00:00",
  "format": "16:9",
  "resolution": "1920x1080",
  "style": {
    "backgroundColor": "#000000",
    "textColor": "#FFFFFF",
    "fontFamily": "Arial",
    "fontSize": 120,
    "fontWeight": "bold",
    "showMilliseconds": false
  },
  "audio": {
    "enabled": false,
    "tickSoundPath": null
  },
  "title": "New Year Countdown"
}""",
    )
    doc.add_heading("RenderJob", level=2)
    add_code_block(
        doc,
        """{
  "id": "uuid",
  "status": "queued | processing | completed | failed | cancelled",
  "progress": 0,
  "elapsedSeconds": null,
  "outputPath": null,
  "outputSizeMb": null,
  "error": null,
  "createdAt": "ISO8601"
}""",
    )
    doc.add_heading("Validation", level=2)
    for line in [
        "durationSeconds: 1 – 28800 (tối đa 8 giờ)",
        "startTime: regex HH:MM:SS",
        "format: 16:9 | 9:16 | 1:1",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 5. API
    doc.add_heading("5. API Endpoints", level=1)
    add_table(
        doc,
        ["Method", "Path", "Mô tả"],
        [
            ("POST", "/api/render", "Nhận RenderConfig, trả { jobId, estimatedSizeMb, estimatedRenderMinutes }"),
            ("GET", "/api/jobs/{id}", "Trạng thái + progress"),
            ("GET", "/api/jobs/{id}/download", "Stream file MP4 (chỉ khi completed)"),
            ("DELETE", "/api/jobs/{id}", "Cancel job + kill FFmpeg"),
            ("GET", "/api/health", "Health check + FFmpeg version"),
        ],
    )

    # 6. Roadmap
    doc.add_heading("6. Lộ trình 5 Phase", level=1)
    add_table(
        doc,
        ["Phase", "Thời gian", "Deliverable"],
        [
            ("0 — Setup", "0.5 ngày", "Repo structure, gitignore, README, FFmpeg check"),
            ("1 — Backend core", "1–2 ngày", "ASS generator + FFmpeg CLI, test 60s → MP4"),
            ("2 — Backend API", "1 ngày", "FastAPI + job manager + progress + cancel"),
            ("3 — Frontend", "1–2 ngày", "Form + preview + polling + download"),
            ("4 — Polish", "1 ngày", "Validation, ước tính file size, error UX, test 3h"),
        ],
    )
    doc.add_paragraph("Thứ tự thực hiện: Phase 0 → 1 → 2 → 3 → 4 → Bonus (tùy chọn)")

    # 7-11 Prompts
    doc.add_heading("7. Prompts triển khai", level=1)
    doc.add_paragraph(
        "Mỗi phase chạy trong một session Agent mode. Dán prompt tương ứng, review output trước khi sang phase tiếp theo."
    )

    add_prompt_block(
        doc,
        "Phase 0 — Khởi tạo project",
        """Tạo monorepo countdown video generator tại thư mục hiện tại với cấu trúc:

- frontend/ — Vite + React + TypeScript + Tailwind CSS
- backend/ — Python FastAPI
- shared/ — placeholder cho JSON schema

Yêu cầu:
1. Root README mô tả dự án: web app tạo video countdown dài (vài giờ), render local bằng FFmpeg
2. .gitignore: node_modules, __pycache__, output/, temp/, .env, *.mp4
3. backend/requirements.txt: fastapi, uvicorn, pydantic, python-multipart
4. backend/app/config.py: OUTPUT_DIR, TEMP_DIR, MAX_DURATION_SECONDS=28800
5. Script hoặc endpoint GET /api/health kiểm tra ffmpeg có trong PATH không
6. frontend proxy API tới localhost:8000 trong vite.config.ts
7. Không implement logic render yet — chỉ scaffold

Chạy được:
- cd backend && uvicorn app.main:app --reload
- cd frontend && npm run dev""",
    )

    add_prompt_block(
        doc,
        "Phase 1 — ASS Generator + FFmpeg (core)",
        """Implement backend render core cho countdown video (Python + FFmpeg).

Context:
- Video countdown từ startTime (HH:MM:SS) đếm ngược durationSeconds giây
- Mỗi giây hiển thị 1 label thời gian (03:00:00 → 02:59:59 → ... → 00:00:00)
- Dùng ASS subtitle burn vào nền solid color — KHÔNG render frame-by-frame

Tạo các file:

1. backend/app/models/schemas.py
   - Pydantic models: RenderConfig, RenderStyle, RenderJob
   - Validation: durationSeconds 1–28800, startTime regex HH:MM:SS

2. backend/app/services/ass_generator.py
   - class ASSGenerator:
     - generate(config) -> str (path to .ass file)
     - Sinh ASS v4+ với PlayResX/PlayResY theo resolution
     - Mỗi giây 1 Dialogue event
     - Style: font, size, color (ASS BGR format), alignment center
     - Nếu có title: event riêng hiển thị suốt video phía trên

3. backend/app/services/ffmpeg_runner.py
   - class FFmpegRunner:
     - build_command(ass_path, config, output_path) -> list[str]
     - Nền: lavfi color=c={bg}:s={resolution}:d={duration}
     - Filter: ass='{ass_path}'
     - Encode: libx264, preset veryfast, crf 23
     - run() async subprocess, parse stderr cho progress (time=...)
     - cancel() kill process

4. CLI test script backend/scripts/test_render.py
   - Render countdown 60 giây từ 00:01:00
   - Output: temp/test_60s.mp4

Yêu cầu kỹ thuật:
- Hàm format_time(seconds) -> "HH:MM:SS"
- ASS timestamps: H:MM:SS.cc
- Unit test cho ass_generator: số events = durationSeconds

Không cần FastAPI routes yet.""",
    )

    add_prompt_block(
        doc,
        "Phase 2 — FastAPI + Job Manager",
        """Implement FastAPI API layer cho countdown video backend.

1. backend/app/services/job_manager.py
   - In-memory dict jobs
   - create_job(config) -> job_id, spawn background task
   - cancel_job(id): kill ffmpeg
   - Background: generate ASS → run FFmpeg → update progress → completed/failed
   - Cleanup temp ASS sau khi xong

2. backend/app/routers/render.py
   - POST /api/render
   - GET /api/jobs/{id}
   - GET /api/jobs/{id}/download
   - DELETE /api/jobs/{id}
   - GET /api/health

3. backend/app/main.py
   - CORS cho localhost:5173
   - Startup: tạo output/ temp/

4. Ước tính:
   - estimatedSizeMb ≈ duration * 4Mbps / 8 / 1024
   - estimatedRenderMinutes ≈ duration / 60 * 0.1

Test bằng curl: POST render 120s → poll → download MP4""",
    )

    add_prompt_block(
        doc,
        "Phase 3 — Frontend UI",
        """Implement frontend cho countdown video generator.

Stack: Vite + React + TypeScript + Tailwind.

1. src/hooks/useCountdown.ts — logic đếm ngược HH:MM:SS
2. src/components/ConfigForm.tsx — duration presets, start time, format, resolution, style
3. src/components/PreviewPlayer.tsx — preview 10–15s, CSS theo aspect ratio
4. src/components/RenderProgress.tsx — poll progress, cancel, download
5. src/hooks/useRenderJob.ts — submitRender, pollJob, cancelJob
6. App.tsx — layout: Form trái, Preview phải, Progress dưới

API base: VITE_API_URL default http://localhost:8000
Dark theme, responsive.""",
    )

    add_prompt_block(
        doc,
        "Phase 4 — Polish & Hardening",
        """Hoàn thiện countdown video generator.

Backend:
- Giới hạn 1 render job concurrent
- Log structured
- Optional audio tick mỗi giây
- Test integration render 3600s (1h)

Frontend:
- Disable submit khi invalid
- Confirm dialog trước render >1h
- Persist config trong localStorage
- Error banner / toast

shared/config.schema.json — JSON Schema
Root README: prerequisites, install, curl example, known limits
Manual test checklist""",
    )

    # 12. Bonus prompts
    doc.add_heading("8. Prompt Bonus (tùy chọn sau MVP)", level=1)
    bonuses = [
        ("Audio tick mỗi giây", "Thêm tick sound: loop tick.wav mỗi giây, merge với -shortest. Cập nhật RenderConfig.audio và frontend toggle."),
        ("Background gradient / image", "backgroundType = solid | gradient | image. Cập nhật ffmpeg_runner và frontend upload."),
        ("Countdown đến datetime cụ thể", "Thêm targetDateTime ISO. Backend tính durationSeconds = target - now. Frontend datetime picker."),
        ("9:16 cho TikTok/Reels", "Format 1080x1920: ASS PlayRes, font scale, preview aspect-ratio 9/16. Test render 60s."),
    ]
    for title, desc in bonuses:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # 13. Checklist
    doc.add_heading("9. Checklist MVP", level=1)
    checklist = [
        "Preview 10s khớp style với video output",
        "Render 60s → download MP4, số đếm đúng từng giây",
        "Render 1h chạy local, progress cập nhật",
        "Cancel job dừng FFmpeg",
        "Validation max duration, invalid time",
        "README đủ để người mới chạy được",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}")

    # 14. Prerequisites
    doc.add_heading("10. Prerequisites", level=1)
    for line in [
        "Python 3.11+",
        "Node.js 20+",
        "FFmpeg (brew install ffmpeg trên macOS)",
        "SSD đủ dung lượng (1h 1080p ≈ 1.5–2.5 GB)",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 15. Notes
    doc.add_heading("11. Ghi chú sử dụng prompt", level=1)
    for line in [
        "Một phase = một session Agent mode",
        "Sau Phase 1: chạy test_render.py trước khi làm Frontend",
        "Nếu contract thay đổi, dán lại RenderConfig vào đầu prompt phase tiếp theo",
        "Local server phù hợp: không lo chi phí cloud, render lâu OK",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
